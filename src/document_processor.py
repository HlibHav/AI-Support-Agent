import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import docx
from docx import Document
import pandas as pd
from datetime import datetime

class DocumentProcessor:
    def __init__(self, knowledge_path: str = "Knowledge"):
        """
        Initialize the document processor for .docx files.
        
        Args:
            knowledge_path: Path to the knowledge directory containing .docx files
        """
        self.knowledge_path = Path(knowledge_path)
        self.documents = []
        self.logger = logging.getLogger(__name__)
        
    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Extract text content from a .docx file.
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            return '\n'.join(text_content)
            
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def get_document_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing document metadata
        """
        stat = file_path.stat()
        relative_path = file_path.relative_to(self.knowledge_path)
        
        # Create category based on file path
        category = "General"
        if "Меню" in str(file_path):
            if "GDS" in str(file_path):
                category = "GDS"
            elif "Заказы" in str(file_path):
                category = "Orders"
            elif "Финансы" in str(file_path):
                category = "Finance"
            elif "Прайсер" in str(file_path):
                category = "Pricing"
            elif "Сайты" in str(file_path):
                category = "Websites"
            elif "Маркетинг" in str(file_path):
                category = "Marketing"
            elif "Справочник" in str(file_path):
                category = "Reference"
            elif "Настройки" in str(file_path):
                category = "Settings"
            elif "Квоты" in str(file_path):
                category = "Quotas"
            elif "Каталог" in str(file_path):
                category = "Catalog"
        elif "F.A.Q" in str(file_path):
            category = "FAQ"
        elif "Разное" in str(file_path):
            category = "Miscellaneous"
        elif "отельный контракт" in str(file_path):
            category = "Hotel Contracts"
        elif "Матчинг" in str(file_path):
            category = "Matching"
        
        return {
            'filename': file_path.name,
            'relative_path': str(relative_path),
            'category': category,
            'size_bytes': stat.st_size,
            'modified_date': datetime.fromtimestamp(stat.st_mtime),
            'file_extension': file_path.suffix.lower()
        }
    
    def process_all_documents(self) -> List[Dict[str, Any]]:
        """
        Process all .docx files in the knowledge directory.
        
        Returns:
            List of processed documents with text content and metadata
        """
        self.documents = []
        
        # Find all .docx files recursively
        docx_files = list(self.knowledge_path.rglob("*.docx"))
        
        self.logger.info(f"Found {len(docx_files)} .docx files to process")
        
        for file_path in docx_files:
            # Skip temporary files
            if file_path.name.startswith('~$'):
                continue
                
            try:
                # Extract text content
                text_content = self.extract_text_from_docx(file_path)
                
                if not text_content.strip():
                    self.logger.warning(f"No text content extracted from {file_path}")
                    continue
                
                # Get metadata
                metadata = self.get_document_metadata(file_path)
                
                # Create document entry
                document = {
                    'id': f"doc_{len(self.documents) + 1}",
                    'text_content': text_content,
                    'title': file_path.stem,
                    'char_count': len(text_content),
                    'word_count': len(text_content.split()),
                    **metadata
                }
                
                self.documents.append(document)
                
                self.logger.info(f"Processed: {file_path.name} ({len(text_content)} chars)")
                
            except Exception as e:
                self.logger.error(f"Error processing {file_path}: {str(e)}")
                continue
        
        self.logger.info(f"Successfully processed {len(self.documents)} documents")
        return self.documents
    
    def get_documents_by_category(self, category: str) -> List[Dict[str, Any]]:
        """
        Get documents filtered by category.
        
        Args:
            category: Category to filter by
            
        Returns:
            List of documents in the specified category
        """
        return [doc for doc in self.documents if doc['category'] == category]
    
    def get_categories(self) -> List[str]:
        """
        Get all unique categories from processed documents.
        
        Returns:
            List of unique categories
        """
        return sorted(set(doc['category'] for doc in self.documents))
    
    def search_documents(self, query: str, category: Optional[str] = None) -> List[Dict[str, Any]]:
        """
        Simple text search in documents.
        
        Args:
            query: Search query
            category: Optional category filter
            
        Returns:
            List of documents matching the search criteria
        """
        query_lower = query.lower()
        results = []
        
        for doc in self.documents:
            if category and doc['category'] != category:
                continue
            
            # Search in title and text content
            if (query_lower in doc['title'].lower() or 
                query_lower in doc['text_content'].lower()):
                results.append(doc)
        
        return results
    
    def to_dataframe(self) -> pd.DataFrame:
        """
        Convert processed documents to a pandas DataFrame.
        
        Returns:
            DataFrame with document information
        """
        if not self.documents:
            return pd.DataFrame()
        
        return pd.DataFrame(self.documents)
    
    def get_summary_stats(self) -> Dict[str, Any]:
        """
        Get summary statistics about the processed documents.
        
        Returns:
            Dictionary with summary statistics
        """
        if not self.documents:
            return {}
        
        df = self.to_dataframe()
        
        return {
            'total_documents': len(self.documents),
            'total_characters': df['char_count'].sum(),
            'total_words': df['word_count'].sum(),
            'average_document_length': df['char_count'].mean(),
            'categories': self.get_categories(),
            'documents_by_category': df.groupby('category').size().to_dict(),
            'largest_document': df.loc[df['char_count'].idxmax(), 'title'],
            'smallest_document': df.loc[df['char_count'].idxmin(), 'title']
        } 